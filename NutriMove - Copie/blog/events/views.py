from django.views import generic
from django.urls import reverse_lazy
from django.shortcuts import render, get_object_or_404, redirect
from .models import Event, Participant
from .forms import ParticipantForm  # Ensure this form is defined
from django.contrib import messages
import csv
from django.http import HttpResponse
from django.template.loader import render_to_string
import pandas as pd
import io
import xhtml2pdf.pisa as pisa
from docx import Document
from django.template.loader import get_template  # Importation de get_template
from django.db.models import Avg
from users.models import Client  # Import your Client model
from django.views import generic
from django.urls import reverse_lazy
from django import forms


class EventListView(generic.ListView):
    model = Event
    template_name = 'events/event_list.html'  # Ensure this path is correct
    context_object_name = 'events'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        for event in context['events']:
            event.capacity_remaining = event.max_participants - event.participants.count()
        return context



class EventCreateForm(forms.ModelForm):
    class Meta:
        model = Event
        fields = [
            'title', 
            'description', 
            'date', 
            'location', 
            'organizer',
            'max_participants', 
            'price', 
            'image'  
        ]

class EventCreateView(generic.CreateView):
    model = Event
    template_name = 'events/event_form.html'
    form_class = EventCreateForm
    success_url = reverse_lazy('event_list')

    def form_valid(self, form):
        return super().form_valid(form)

class EventDetailView(generic.DetailView):
    model = Event
    template_name = 'events/event_detail.html'

class EventUpdateView(generic.UpdateView):
    model = Event
    template_name = 'events/event_form.html'
    fields = [
        'title', 
        'description', 
        'date', 
        'location', 
        'organizer', 
        'max_participants', 
        'price', 
        'image', 
        'organizer'
        ]  
    success_url = reverse_lazy('event_list')

    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        
        # Pre-fill the date with the current value and make it read-only
        form.fields['date'].initial = self.object.date
        form.fields['date'].widget.attrs['readonly'] = True
        return form

    def form_valid(self, form):
        # Always keep the original date
        form.instance.date = self.object.date

        # Keep the original image if no new image is provided
        if not form.cleaned_data.get('image'):
            form.instance.image = self.object.image
        
        return super().form_valid(form)



class EventDeleteView(generic.DeleteView):
    model = Event
    template_name = 'events/event_confirm_delete.html'
    success_url = reverse_lazy('event_list')

def participant_list(request, event_id):
    event = get_object_or_404(Event, id=event_id)
    participants = Participant.objects.filter(event=event)
    capacity_remaining = event.max_participants - participants.count()

    return render(request, 'participant/participant_list.html', {
        'event': event,
        'participants': participants,
        'capacity_remaining': capacity_remaining,
    })



def participant_create(request, event_id):
    event = get_object_or_404(Event, id=event_id)
    participants = Participant.objects.filter(event=event)
    clients = Client.objects.all()

    # Vérifier la capacité restante
    capacity_remaining = event.max_participants - participants.count()

    # Rediriger si la capacité est atteinte
    if capacity_remaining <= 0:
        return redirect('participant_list', event_id=event.id)

    if request.method == 'POST':
        form = ParticipantForm(request.POST)
        if form.is_valid():
            client = form.cleaned_data['client']
            
            # Vérifier si le client a déjà participé à l'événement
            if participants.filter(client=client).exists():
                messages.error(request, "Ce client a déjà participé à cet événement.")
                return redirect('participant_list', event_id=event.id)
            else:
                if participants.count() < event.max_participants:
                    participant = form.save(commit=False)
                    participant.event = event
                    participant.save()
                    return redirect('participant_list', event_id=event.id)
                else:
                    form.add_error(None, "La capacité maximale de participants a été atteinte.")
    else:
        form = ParticipantForm()

    return render(request, 'participant/participant_create.html', {
        'event': event,
        'form': form,
        'clients': clients,
    })



def participant_update(request, event_id, participant_id):
    event = get_object_or_404(Event, pk=event_id)
    participant = get_object_or_404(Participant, pk=participant_id)

    if request.method == 'POST':
        # Ensure the client ID is passed correctly in the POST data
        client_id = request.POST.get('client')
        if client_id:
            client = get_object_or_404(Client, pk=client_id)
            participant.client = client
        
        form = ParticipantForm(request.POST, instance=participant)
        if form.is_valid():
            form.save()
            return redirect('participant_list', event_id=event_id)
    else:
        form = ParticipantForm(instance=participant)

    return render(request, 'participant/participant_update.html', {
        'participant': participant,
        'form': form,
        'event': event,
    })

def participant_delete(request, event_id, participant_id):
    participant = get_object_or_404(Participant, pk=participant_id)
    if request.method == 'POST':
        participant.delete()
        return redirect('participant_list', event_id=event_id)
    return render(request, 'participant/participant_confirm_delete.html', {'participant': participant})

def event_export(request, format):
    events = Event.objects.all()

    if format == 'pdf':
        # Générer un PDF
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename=events.pdf'

        # Code de génération de PDF
        template_path = 'events/event_pdf_template.html'  # Chemin vers votre template PDF
        template = get_template(template_path)
        context = {'events': events}
        html = template.render(context)

        pisa_status = pisa.CreatePDF(html, dest=response)

        if pisa_status.err:
            return HttpResponse('Une erreur est survenue lors de la génération du PDF.')

        return response

    elif format == 'docs':
        # Générer un fichier Docs
        response = HttpResponse(content_type='application/msword')
        response['Content-Disposition'] = 'attachment; filename=events.doc'

        doc = Document()
        doc.add_heading('Liste des Événements', level=1)

        for event in events:
            doc.add_heading(event.title, level=2)
            doc.add_paragraph(f'Description: {event.description}')
            doc.add_paragraph(f'Date: {event.date.strftime("%d/%m/%Y %H:%M")}')
            doc.add_paragraph(f'Lieu: {event.location}')
            doc.add_paragraph(f'Organisateur: {event.organizer}')
            doc.add_paragraph(f'Participants Maximum: {event.max_participants}')
            doc.add_paragraph(f'Participants Restants: {event.capacity_remaining}')

            doc.add_paragraph(f'Prix: {event.price} €')
            doc.add_paragraph(f'Date de création: {event.created_at}')
            doc.add_paragraph(f'Date de  Modification: {event.updated_at}')

            doc.add_paragraph('---')  # Ligne de séparation entre les événements

        doc.save(response)  # Enregistrer le document dans la réponse
        return response

    elif format == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="events.csv"'

        writer = csv.writer(response)
        writer.writerow(['Titre', 'Description', 'Date', 'Lieu', 'Organisateur', 'Participants Maximum', 'Prix', 'Catégorie'])  # En-tête

        for event in events:
            writer.writerow([
                event.title,
                event.description,
                event.date.strftime('%d/%m/%Y %H:%M'),  # Formatage de la date
                event.location,
                event.organizer,
                event.max_participants,
                event.price,
                event.category
            ])

        return response

    else:
        # Gérer les formats non pris en charge
        return HttpResponse("Format non pris en charge", status=400)
